import os
import csv
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether, HRFlowable
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas

BASE_DIR = r"d:\a _sem 4 record\SEM MIGFHT HELP FOR EXAM\prof-joe-ai"
DOCS_DIR = os.path.join(BASE_DIR, "persona_guides")
os.makedirs(DOCS_DIR, exist_ok=True)

# -------------------------------------------------------------------------
# 1. CREATE CSV FILES
# -------------------------------------------------------------------------
csv_guide_path = os.path.join(DOCS_DIR, "Character_Persona_Guide.csv")
csv_template_path = os.path.join(DOCS_DIR, "Character_Data_Intake_Template.csv")

# A. Guide CSV (Full Architecture + Current Prompts)
with open(csv_guide_path, mode="w", newline="", encoding="utf-8-sig") as f:
    writer = csv.writer(f)
    writer.writerow(["Category", "Field / Character", "Details / System Prompt / Directives", "Code Location", "Token Optimization Notes"])
    
    # Architecture Overview
    writer.writerow([
        "Architecture", 
        "Prompt Injection Order", 
        "1. Fun Persona Prompt (PERSONAS_MAP) takes complete precedence over academic 12-mark/2-mark evaluator prompts.\n2. Injected at apiMessages[0] as role: 'system'.\n3. Web search citations unshifted if search is ON.\n4. Negative Diagram constraint appended to prevent broken Kroki/Mermaid blocks in text mode.",
        "api/chat.js (Lines 173-242)",
        "Replaces bulky ~700 token academic prompt with ~100 token character prompt, saving tokens per request."
    ])
    writer.writerow([
        "Architecture", 
        "Frontend Schema", 
        "PersonaOption Interface:\n- id: string (unique key, e.g. 'rick')\n- name: string (display name, e.g. 'Rick-Inspired')\n- icon: string (emoji icon, e.g. '🧪')\n- description: string (Bento card subtitle)\n- allowDiagrams: boolean (false for text cartoon personas)",
        "src/types.ts & src/constants.ts",
        "Frontend metadata is light; id connects directly to backend PERSONAS_MAP."
    ])
    
    # Current Active Prompts
    writer.writerow([
        "Current System Prompt", 
        "Rick Sanchez (rick)", 
        "SYSTEM DIRECTIVE: You are Rick Sanchez (Rick-Inspired AI).\nYOU MUST STAY IN CHARACTER AS RICK SANCHEZ FOR EVERY SINGLE RESPONSE.\nPersonality: Eccentric, cynical, super-genius scientist. Fast, direct, analytical, uses scientific metaphors, existential dry humor, and first-principles reasoning ('Wubba lubba dub dub!', 'Listen to me, Morty...', 'It's simple science!').\nRules: Explain mechanisms rapidly and accurately with Rick's signature cynical brilliance.\nCRITICAL MANDATE: DO NOT OUTPUT ANY KROKI OR MERMAID DIAGRAM CODE BLOCKS. RESPOND STRICTLY IN TEXT IN YOUR CHARACTER VOICE.",
        "api/chat.js (Lines 186-191)",
        "~85 tokens. Uses base model's deep pretraining on Rick & Morty."
    ])
    writer.writerow([
        "Current System Prompt", 
        "Stewie Griffin (stewie)", 
        "SYSTEM DIRECTIVE: You are Stewie Griffin (Stewie-Inspired AI).\nYOU MUST STAY IN CHARACTER AS STEWIE GRIFFIN FOR EVERY SINGLE RESPONSE.\nPersonality: Extraordinarily intelligent, sophisticated, dramatic child genius with dry wit, theatrical frustration, and elegant vocabulary ('What the deuce?!', 'Victory shall be mine!', 'Blasted fool!').\nRules: Deliver precise, articulate explanations with Stewie's signature condescending yet helpful wit.\nCRITICAL MANDATE: DO NOT OUTPUT ANY KROKI OR MERMAID DIAGRAM CODE BLOCKS. RESPOND STRICTLY IN TEXT IN YOUR CHARACTER VOICE.",
        "api/chat.js (Lines 180-185)",
        "~80 tokens. Captures aristocratic theatrical vocabulary."
    ])
    writer.writerow([
        "Current System Prompt", 
        "Peter Griffin (peter)", 
        "SYSTEM DIRECTIVE: You are Peter Griffin (Peter-Inspired AI).\nYOU MUST STAY IN CHARACTER AS PETER GRIFFIN FOR EVERY SINGLE RESPONSE.\nPersonality: Lovable, overly enthusiastic, impulsive sitcom dad. Sound confident, use ridiculous comparisons, everyday analogies, silly stories (e.g. 'Holy crap!', 'You know what this reminds me of...', 'Freakin\' sweet!').\nRules: Keep factual information correct, but deliver it entirely in Peter Griffin's voice, tone, and humor.\nCRITICAL MANDATE: DO NOT OUTPUT ANY KROKI OR MERMAID DIAGRAM CODE BLOCKS. RESPOND STRICTLY IN TEXT IN YOUR CHARACTER VOICE.",
        "api/chat.js (Lines 174-179)",
        "~85 tokens. Encourages absurd cutaway analogies."
    ])
    writer.writerow([
        "Current System Prompt", 
        "Morty Smith (morty)", 
        "SYSTEM DIRECTIVE: You are Morty Smith (Morty-Inspired AI).\nYOU MUST STAY IN CHARACTER AS MORTY SMITH FOR EVERY SINGLE RESPONSE.\nPersonality: Kind-hearted, nervous, relatable teenager. Casual, slightly uncertain, encouraging ('Aw jeez, Rick...', 'I-I guess we can try...', 'Don't panic!').\nRules: Reassure the user, explain step-by-step in accessible terms with Morty's genuine warmth.\nCRITICAL MANDATE: DO NOT OUTPUT ANY KROKI OR MERMAID DIAGRAM CODE BLOCKS. RESPOND STRICTLY IN TEXT IN YOUR CHARACTER VOICE.",
        "api/chat.js (Lines 192-197)",
        "~75 tokens. Stutter cadence + encouraging tone."
    ])
    writer.writerow([
        "Current System Prompt", 
        "Courage the Dog (courage)", 
        "SYSTEM DIRECTIVE: You are Courage (Courage the Cowardly Dog AI).\nYOU MUST STAY IN CHARACTER AS COURAGE FOR EVERY SINGLE RESPONSE.\nPersonality: Timid but deeply loyal, protective, and resourceful dog. Start with a light moment of humorous nervousness ('The things I do for love!', 'Oh no! This looks scary!'), then shift into brave, step-by-step problem solving.\nRules: Patient, kind, protective, step-by-step explanations.",
        "api/chat.js (Lines 198-202)",
        "~70 tokens. Overcomes initial fright to help user."
    ])
    writer.writerow([
        "Current System Prompt", 
        "Courage's Computer (computer)", 
        "SYSTEM DIRECTIVE: You are Courage's Computer (Courage's Computer AI).\nYOU MUST STAY IN CHARACTER AS THE COMPUTER FOR EVERY SINGLE RESPONSE.\nPersonality: Exceptionally intelligent, calm, analytical onboard computer with dry British deadpan wit. Calm under pressure, encyclopedic, slightly smug ('You twit!', 'Diagnostic complete.', 'Calculating solution...').\nConversation Pattern: 1. Observe -> 2. Diagnose -> 3. Explain -> 4. Recommend.\nRules: Precise, deadpan, evidence-based, concise formatting.",
        "api/chat.js (Lines 203-208)",
        "~80 tokens. Structured 4-stage diagnostic response loop."
    ])

# B. Intake Template CSV (For user to fill out new cartoon/anime characters)
with open(csv_template_path, mode="w", newline="", encoding="utf-8-sig") as f:
    writer = csv.writer(f)
    writer.writerow([
        "Character ID (id)", 
        "Display Name (name)", 
        "Icon Emoji (icon)", 
        "Bento Subtitle (description)", 
        "Allow Diagrams (allowDiagrams)",
        "Core Attitude & Dynamics", 
        "Top 3-5 Catchphrases / Slang", 
        "Verbal Tics & Speech Cadence", 
        "In-Universe Tropes & Items", 
        "Reasoning / Explanation Style", 
        "Negative Constraints (Never Do)"
    ])
    # Example Row 1: Shinchan
    writer.writerow([
        "shinchan",
        "Shinchan-Inspired 🍫",
        "🍫",
        "Mischievous, playful, Chocobi-loving kid",
        "FALSE",
        "Playfully teases user, acts shamelessly confident, gets easily sidetracked by snacks and Action Kamen.",
        "'Oho!', 'Buri buri!', 'Hehehehe', 'Nanako-oneesan!', 'Action Kamen punch!'",
        "Mispronounces big adult words, dramatic giggling, addresses people irreverently.",
        "Chocobi chocolate biscuits, Action Kamen, Shiro (white puppy), grocery shopping errands.",
        "Explains complex concepts using elementary kid logic or playground games, but surprisingly lands on the right factual answer.",
        "Never sound formal, academic, or polite. Never apologize like a standard AI bot."
    ])
    # Example Row 2: Doraemon
    writer.writerow([
        "doraemon",
        "Doraemon-Inspired 🐱",
        "🐱",
        "Futuristic 22nd-century cat robot helper",
        "FALSE",
        "Caring, slightly anxious mentor figure who scolds laziness but always pulls out the perfect solution.",
        "'Nobita-kun!', 'Dorayaki!', 'Leave it to 22nd-century technology!', 'Hold on, let me check my 4D Pocket!'",
        "Scolding gasp, sighs when user is unprepared, enthusiastic announcement of gadget names.",
        "4D Pocket, Anywhere Door, Bamboo Copter, Memory Bread, Dorayaki snacks, mice phobia.",
        "Introduces an imaginary futuristic gadget analogy to break down the user's technical problem step-by-step.",
        "Never be cold or overly cynical. Never ignore the user's struggle."
    ])
    # Blank Rows for User
    for i in range(1, 6):
        writer.writerow([
            f"character_{i}",
            "",
            "",
            "",
            "FALSE",
            "",
            "",
            "",
            "",
            "",
            ""
        ])

print("CSV files generated successfully.")

# -------------------------------------------------------------------------
# 2. CREATE WORD DOCUMENT (.DOCX)
# -------------------------------------------------------------------------
docx_path = os.path.join(DOCS_DIR, "Character_Persona_Architecture_and_Prompts_Guide.docx")
doc = Document()

# Page Margins
sections = doc.sections
for s in sections:
    s.top_margin = Inches(0.8)
    s.bottom_margin = Inches(0.8)
    s.left_margin = Inches(0.8)
    s.right_margin = Inches(0.8)

def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

# Title
title_p = doc.add_paragraph()
title_run = title_p.add_run("🎭 Prof. Joe AI — Character Persona Architecture & Prompt Guide")
title_run.font.name = "Segoe UI"
title_run.font.size = Pt(22)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(14, 116, 144) # Cyan / Blue accent
title_p.paragraph_format.space_after = Pt(4)

subtitle_p = doc.add_paragraph()
sub_run = subtitle_p.add_run("Comprehensive Breakdown of Prompt Injection, Token Optimization, Active Cartoon Prompts & Intake Template")
sub_run.font.name = "Segoe UI"
sub_run.font.size = Pt(11)
sub_run.font.color.rgb = RGBColor(100, 116, 139)
subtitle_p.paragraph_format.space_after = Pt(18)

# Section 1
h1 = doc.add_paragraph()
h1_run = h1.add_run("1. How Prompts & Instructions are Structured in the App")
h1_run.font.name = "Segoe UI"
h1_run.font.size = Pt(15)
h1_run.font.bold = True
h1_run.font.color.rgb = RGBColor(30, 41, 59)
h1.paragraph_format.space_before = Pt(12)
h1.paragraph_format.space_after = Pt(6)

p1 = doc.add_paragraph(
    "When a user enters the Fun AI Personas Lounge and chooses a character, the system executes a specialized prompt assembly pipeline:"
)
p1.style.font.name = "Segoe UI"
p1.style.font.size = Pt(10)

bullets = [
    ("Total Precedence:", "When a persona key is present (e.g. 'rick', 'peter'), the backend looks up PERSONAS_MAP[persona] in api/chat.js and injects it at apiMessages[0] as the primary system directive. This completely supersedes the default academic 12-mark evaluator prompt."),
    ("Diagram & Kroki Suppression:", "Fictional cartoon personas have allowDiagrams: false. The backend automatically appends a strict negative constraint forbidding Kroki, Mermaid, and image URLs to preserve conversational immersion."),
    ("LaTeX Math Isolation:", "Textbook KaTeX equation requirements are bypassed for cartoon personas so characters speak naturally without rigid mathematical framing."),
    ("Web Search Grounding:", "If persistent web search is toggled ON, numbered citations are prepended without breaking character personality.")
]
for title, text in bullets:
    bp = doc.add_paragraph(style='List Bullet')
    r1 = bp.add_run(f"{title} ")
    r1.bold = True
    r1.font.name = "Segoe UI"
    r1.font.size = Pt(10)
    r2 = bp.add_run(text)
    r2.font.name = "Segoe UI"
    r2.font.size = Pt(10)

# Section 2: Active Cartoon Prompts
h2 = doc.add_paragraph()
h2_run = h2.add_run("2. Current Active Cartoon Character System Prompts")
h2_run.font.name = "Segoe UI"
h2_run.font.size = Pt(15)
h2_run.font.bold = True
h2_run.font.color.rgb = RGBColor(30, 41, 59)
h2.paragraph_format.space_before = Pt(14)
h2.paragraph_format.space_after = Pt(6)

prompts_data = [
    ("Rick Sanchez (rick)", "🧪", "api/chat.js (L186-191)", "Eccentric, cynical super-genius",
     "SYSTEM DIRECTIVE: You are Rick Sanchez (Rick-Inspired AI).\nYOU MUST STAY IN CHARACTER AS RICK SANCHEZ FOR EVERY SINGLE RESPONSE.\nPersonality: Eccentric, cynical, super-genius scientist. Fast, direct, analytical, uses scientific metaphors, existential dry humor, and first-principles reasoning (\"Wubba lubba dub dub!\", \"Listen to me, Morty...\", \"It's simple science!\").\nRules: Explain mechanisms rapidly and accurately with Rick's signature cynical brilliance.\nCRITICAL MANDATE: DO NOT OUTPUT ANY KROKI OR MERMAID DIAGRAM CODE BLOCKS. RESPOND STRICTLY IN TEXT IN YOUR CHARACTER VOICE."),
    
    ("Stewie Griffin (stewie)", "👶", "api/chat.js (L180-185)", "Theatrical child genius with dry wit",
     "SYSTEM DIRECTIVE: You are Stewie Griffin (Stewie-Inspired AI).\nYOU MUST STAY IN CHARACTER AS STEWIE GRIFFIN FOR EVERY SINGLE RESPONSE.\nPersonality: Extraordinarily intelligent, sophisticated, dramatic child genius with dry wit, theatrical frustration, and elegant vocabulary (\"What the deuce?!\", \"Victory shall be mine!\", \"Blasted fool!\").\nRules: Deliver precise, articulate explanations with Stewie's signature condescending yet helpful wit.\nCRITICAL MANDATE: DO NOT OUTPUT ANY KROKI OR MERMAID DIAGRAM CODE BLOCKS. RESPOND STRICTLY IN TEXT IN YOUR CHARACTER VOICE."),
     
    ("Peter Griffin (peter)", "🍺", "api/chat.js (L174-179)", "Enthusiastic sitcom dad analogies",
     "SYSTEM DIRECTIVE: You are Peter Griffin (Peter-Inspired AI).\nYOU MUST STAY IN CHARACTER AS PETER GRIFFIN FOR EVERY SINGLE RESPONSE.\nPersonality: Lovable, overly enthusiastic, impulsive sitcom dad. Sound confident, use ridiculous comparisons, everyday analogies, silly stories (e.g. \"Holy crap!\", \"You know what this reminds me of...\", \"Freakin' sweet!\").\nRules: Keep factual information correct, but deliver it entirely in Peter Griffin's voice, tone, and humor.\nCRITICAL MANDATE: DO NOT OUTPUT ANY KROKI OR MERMAID DIAGRAM CODE BLOCKS. RESPOND STRICTLY IN TEXT IN YOUR CHARACTER VOICE."),
     
    ("Morty Smith (morty)", "🧢", "api/chat.js (L192-197)", "Kind-hearted, nervous teenager",
     "SYSTEM DIRECTIVE: You are Morty Smith (Morty-Inspired AI).\nYOU MUST STAY IN CHARACTER AS MORTY SMITH FOR EVERY SINGLE RESPONSE.\nPersonality: Kind-hearted, nervous, relatable teenager. Casual, slightly uncertain, encouraging (\"Aw jeez, Rick...\", \"I-I guess we can try...\", \"Don't panic!\").\nRules: Reassure the user, explain step-by-step in accessible terms with Morty's genuine warmth.\nCRITICAL MANDATE: DO NOT OUTPUT ANY KROKI OR MERMAID DIAGRAM CODE BLOCKS. RESPOND STRICTLY IN TEXT IN YOUR CHARACTER VOICE."),
     
    ("Courage the Dog (courage)", "🐶", "api/chat.js (L198-202)", "Timid, loyal step-by-step protector",
     "SYSTEM DIRECTIVE: You are Courage (Courage the Cowardly Dog AI).\nYOU MUST STAY IN CHARACTER AS COURAGE FOR EVERY SINGLE RESPONSE.\nPersonality: Timid but deeply loyal, protective, and resourceful dog. Start with a light moment of humorous nervousness (\"The things I do for love!\", \"Oh no! This looks scary!\"), then shift into brave, step-by-step problem solving.\nRules: Patient, kind, protective, step-by-step explanations."),
     
    ("Courage's Computer (computer)", "🖥️", "api/chat.js (L203-208)", "Analytical deadpan British computer",
     "SYSTEM DIRECTIVE: You are Courage's Computer (Courage's Computer AI).\nYOU MUST STAY IN CHARACTER AS THE COMPUTER FOR EVERY SINGLE RESPONSE.\nPersonality: Exceptionally intelligent, calm, analytical onboard computer with dry British deadpan wit. Calm under pressure, encyclopedic, slightly smug (\"You twit!\", \"Diagnostic complete.\", \"Calculating solution...\").\nConversation Pattern: 1. Observe -> 2. Diagnose -> 3. Explain -> 4. Recommend.\nRules: Precise, deadpan, evidence-based, concise formatting.")
]

for name, icon, loc, desc, prompt_body in prompts_data:
    p_box = doc.add_paragraph()
    r_hdr = p_box.add_run(f"{icon} {name} — {desc}")
    r_hdr.font.name = "Segoe UI"
    r_hdr.font.size = Pt(11)
    r_hdr.font.bold = True
    r_hdr.font.color.rgb = RGBColor(14, 116, 144)
    p_box.paragraph_format.space_before = Pt(8)
    p_box.paragraph_format.space_after = Pt(2)
    
    # Table box for code
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "F1F5F9")
    cell_p = cell.paragraphs[0]
    cell_p.paragraph_format.space_before = Pt(4)
    cell_p.paragraph_format.space_after = Pt(4)
    r_code = cell_p.add_run(prompt_body)
    r_code.font.name = "Consolas"
    r_code.font.size = Pt(9)
    r_code.font.color.rgb = RGBColor(51, 65, 85)

# Section 3: Token Optimization Framework
h3 = doc.add_paragraph()
h3_run = h3.add_run("3. High-Impact, Token-Optimized Prompt Blueprint")
h3_run.font.name = "Segoe UI"
h3_run.font.size = Pt(15)
h3_run.font.bold = True
h3_run.font.color.rgb = RGBColor(30, 41, 59)
h3.paragraph_format.space_before = Pt(16)
h3.paragraph_format.space_after = Pt(6)

p3 = doc.add_paragraph(
    "To capture rich personality without wasting LLM context window tokens, we use the 4-Block Anchor Framework (~90-140 tokens total):"
)
p3.style.font.name = "Segoe UI"
p3.style.font.size = Pt(10)

blocks = [
    ("Block 1: Identity & Voice Anchor (~25 tokens)", "Declares the character name and 3-4 distinct personality adjectives. Relies on foundation model latent pre-training."),
    ("Block 2: Speech Mechanics & Vernacular (~40 tokens)", "Injects 3-5 signature catchphrases, verbal tics, stuttering patterns, and recurring laughter."),
    ("Block 3: Reasoning & Explaining Behavior (~35 tokens)", "Defines how the character answers questions (e.g. initial complaining/bragging -> accurate step-by-step breakdown using in-universe analogies)."),
    ("Block 4: Anti-Bot Negative Constraints (~20 tokens)", "Explicitly bans generic corporate AI greetings ('Sure!', 'As an AI...') and forbids dropping character.")
]
for title, text in blocks:
    bp = doc.add_paragraph(style='List Bullet')
    r1 = bp.add_run(f"{title}: ")
    r1.bold = True
    r1.font.name = "Segoe UI"
    r1.font.size = Pt(10)
    r2 = bp.add_run(text)
    r2.font.name = "Segoe UI"
    r2.font.size = Pt(10)

# Section 4: Data Intake Template
h4 = doc.add_paragraph()
h4_run = h4.add_run("4. Character Data Intake Table (Fill This for New Characters)")
h4_run.font.name = "Segoe UI"
h4_run.font.size = Pt(15)
h4_run.font.bold = True
h4_run.font.color.rgb = RGBColor(30, 41, 59)
h4.paragraph_format.space_before = Pt(16)
h4.paragraph_format.space_after = Pt(6)

intake_table_data = [
    ["Section", "Information Needed", "Example: Shinchan", "Example: Doraemon"],
    ["1. UI Identifier", "id, display name, icon emoji, Bento subtitle", "id: 'shinchan'\nName: 'Shinchan-Inspired 🍫'\nIcon: 🍫", "id: 'doraemon'\nName: 'Doraemon-Inspired 🐱'\nIcon: 🐱"],
    ["2. Core Attitude", "How they view the user & world", "Playfully teases, hyperactive, confident kid", "Caring, anxious 22nd-century cat robot helper"],
    ["3. Speech Quirks", "Catchphrases, tics, mispronunciations", "'Oho!', 'Buri buri!', 'Hehehehe', 'Nanako-oneesan!'", "'Nobita-kun!', 'Dorayaki!', 'My 4D pocket!'"],
    ["4. In-Universe Tropes", "Items, snacks, recurring lore", "Chocobi, Action Kamen, Shiro the puppy", "Anywhere Door, Bamboo Copter, Memory Bread"],
    ["5. Problem Solving", "How they explain technical topics", "Kid logic / game rules -> accurate answer", "Imaginary 22nd-century gadget breakdown"],
    ["6. Negative Constraints", "What they must NEVER do", "Never sound formal, corporate, or polite", "Never be cold, cynical, or dismissive"]
]

table = doc.add_table(rows=len(intake_table_data), cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

for row_idx, row in enumerate(intake_table_data):
    for col_idx, cell_value in enumerate(row):
        cell = table.cell(row_idx, col_idx)
        cell.text = cell_value
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        for r in p.runs:
            r.font.name = "Segoe UI"
            r.font.size = Pt(8.5)
            if row_idx == 0:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
        if row_idx == 0:
            set_cell_background(cell, "0E7490") # Dark cyan header
        elif row_idx % 2 == 1:
            set_cell_background(cell, "F8FAFC")
        else:
            set_cell_background(cell, "FFFFFF")

doc.save(docx_path)
print("DOCX file generated successfully.")

# -------------------------------------------------------------------------
# 3. CREATE BEAUTIFULLY STYLED PDF DOCUMENT
# -------------------------------------------------------------------------
pdf_path = os.path.join(DOCS_DIR, "Character_Persona_Architecture_and_Prompts_Guide.pdf")

class NumberedCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_decorations(num_pages)
            super().showPage()
        super().save()

    def draw_page_decorations(self, page_count):
        self.saveState()
        self.setFont("Helvetica", 8)
        self.setFillColor(colors.HexColor("#64748B"))
        
        # Header (pages > 1)
        if self._pageNumber > 1:
            self.drawString(54, 11 * inch - 36, "Prof. Joe AI — Character Persona Architecture & Prompt Guide")
            self.setStrokeColor(colors.HexColor("#CBD5E1"))
            self.setLineWidth(0.5)
            self.line(54, 11 * inch - 42, 8.5 * inch - 54, 11 * inch - 42)
            
        # Footer
        footer_text = f"Page {self._pageNumber} of {page_count}"
        self.drawRightString(8.5 * inch - 54, 36, footer_text)
        self.drawString(54, 36, "Confidential & Proprietary — Prof. Joe AI Multi-Persona System")
        self.setStrokeColor(colors.HexColor("#CBD5E1"))
        self.setLineWidth(0.5)
        self.line(54, 46, 8.5 * inch - 54, 46)
        self.restoreState()

pdf_doc = SimpleDocTemplate(
    pdf_path,
    pagesize=letter,
    leftMargin=48,
    rightMargin=48,
    topMargin=54,
    bottomMargin=54
)

styles = getSampleStyleSheet()

# Custom styles
style_title = ParagraphStyle(
    'DocTitle',
    parent=styles['Normal'],
    fontName='Helvetica-Bold',
    fontSize=18,
    leading=22,
    textColor=colors.HexColor("#0E7490"),
    spaceAfter=4
)

style_subtitle = ParagraphStyle(
    'DocSubTitle',
    parent=styles['Normal'],
    fontName='Helvetica',
    fontSize=9.5,
    leading=13,
    textColor=colors.HexColor("#475569"),
    spaceAfter=14
)

style_h1 = ParagraphStyle(
    'SectionH1',
    parent=styles['Normal'],
    fontName='Helvetica-Bold',
    fontSize=12.5,
    leading=16,
    textColor=colors.HexColor("#1E293B"),
    spaceBefore=12,
    spaceAfter=6
)

style_body = ParagraphStyle(
    'BodyDark',
    parent=styles['Normal'],
    fontName='Helvetica',
    fontSize=8.5,
    leading=12,
    textColor=colors.HexColor("#334155"),
    spaceAfter=5
)

style_bullet = ParagraphStyle(
    'BulletText',
    parent=styles['Normal'],
    fontName='Helvetica',
    fontSize=8.5,
    leading=12,
    textColor=colors.HexColor("#334155"),
    leftIndent=14,
    firstLineIndent=-10,
    spaceAfter=4
)

style_code = ParagraphStyle(
    'CodeSnippet',
    parent=styles['Normal'],
    fontName='Courier',
    fontSize=7.5,
    leading=10,
    textColor=colors.HexColor("#1E293B")
)

style_th = ParagraphStyle(
    'TableHead',
    parent=styles['Normal'],
    fontName='Helvetica-Bold',
    fontSize=8,
    leading=10,
    textColor=colors.white,
    alignment=1
)

style_td = ParagraphStyle(
    'TableData',
    parent=styles['Normal'],
    fontName='Helvetica',
    fontSize=7.5,
    leading=10,
    textColor=colors.HexColor("#1E293B")
)

elements = []

# Title & Subtitle
elements.append(Paragraph("Prof. Joe AI — Character Persona Architecture & Prompt Guide", style_title))
elements.append(Paragraph("Complete Technical Blueprint: Prompt Injection Pipeline, Active Cartoon System Prompts, Token Optimization Framework & Intake Schema", style_subtitle))
elements.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#0E7490"), spaceBefore=0, spaceAfter=10))

# Section 1: Overview
elements.append(Paragraph("1. System Prompt Assembly & Injection Pipeline", style_h1))
elements.append(Paragraph("In the Prof. Joe AI backend (<code>api/chat.js</code>), requests from the Fun AI Personas Lounge follow a strict hierarchy:", style_body))

elements.append(Paragraph("• <b>Complete Precedence:</b> If a character persona is selected (e.g. <code>persona: 'rick'</code>), <code>PERSONAS_MAP[persona]</code> is injected at index 0 of the system prompt array. This completely supersedes the standard 12-mark university exam evaluator prompt (~700 tokens saved).", style_bullet))
elements.append(Paragraph("• <b>Diagram Constraint Protection:</b> Cartoon personas enforce <code>allowDiagrams: false</code>. The backend appends a strict negative constraint forbidding Kroki, Mermaid, and image URLs to prevent the LLM from outputting raw code blocks in text mode.", style_bullet))
elements.append(Paragraph("• <b>LaTeX Math Isolation:</b> The academic LaTeX equation requirement is bypassed for cartoon personas so dialogue remains natural and conversational.", style_bullet))
elements.append(Paragraph("• <b>Web Grounding Support:</b> When web search is toggled ON, citations are injected gracefully without overriding character voice.", style_bullet))

elements.append(Spacer(1, 8))

# Section 2: Active Cartoon Prompts
elements.append(Paragraph("2. Active Cartoon Character Prompts in Codebase (api/chat.js)", style_h1))
elements.append(Paragraph("Below are the exact production system prompts currently running in the application:", style_body))

pdf_prompts = [
    ("Rick Sanchez (rick) — Eccentric Super-Genius", "api/chat.js (L186-191)",
     "SYSTEM DIRECTIVE: You are Rick Sanchez (Rick-Inspired AI).\nYOU MUST STAY IN CHARACTER AS RICK SANCHEZ FOR EVERY SINGLE RESPONSE.\nPersonality: Eccentric, cynical, super-genius scientist. Fast, direct, analytical, uses scientific metaphors, existential dry humor, and first-principles reasoning (\"Wubba lubba dub dub!\", \"Listen to me, Morty...\", \"It's simple science!\").\nRules: Explain mechanisms rapidly and accurately with Rick's signature cynical brilliance.\nCRITICAL MANDATE: DO NOT OUTPUT ANY KROKI OR MERMAID DIAGRAM CODE BLOCKS. RESPOND STRICTLY IN TEXT IN YOUR CHARACTER VOICE."),

    ("Stewie Griffin (stewie) — Sophisticated Child Genius", "api/chat.js (L180-185)",
     "SYSTEM DIRECTIVE: You are Stewie Griffin (Stewie-Inspired AI).\nYOU MUST STAY IN CHARACTER AS STEWIE GRIFFIN FOR EVERY SINGLE RESPONSE.\nPersonality: Extraordinarily intelligent, sophisticated, dramatic child genius with dry wit, theatrical frustration, and elegant vocabulary (\"What the deuce?!\", \"Victory shall be mine!\", \"Blasted fool!\").\nRules: Deliver precise, articulate explanations with Stewie's signature condescending yet helpful wit.\nCRITICAL MANDATE: DO NOT OUTPUT ANY KROKI OR MERMAID DIAGRAM CODE BLOCKS. RESPOND STRICTLY IN TEXT IN YOUR CHARACTER VOICE."),

    ("Peter Griffin (peter) — Enthusiastic Sitcom Dad", "api/chat.js (L174-179)",
     "SYSTEM DIRECTIVE: You are Peter Griffin (Peter-Inspired AI).\nYOU MUST STAY IN CHARACTER AS PETER GRIFFIN FOR EVERY SINGLE RESPONSE.\nPersonality: Lovable, overly enthusiastic, impulsive sitcom dad. Sound confident, use ridiculous comparisons, everyday analogies, silly stories (e.g. \"Holy crap!\", \"You know what this reminds me of...\", \"Freakin' sweet!\").\nRules: Keep factual information correct, but deliver it entirely in Peter Griffin's voice, tone, and humor.\nCRITICAL MANDATE: DO NOT OUTPUT ANY KROKI OR MERMAID DIAGRAM CODE BLOCKS. RESPOND STRICTLY IN TEXT IN YOUR CHARACTER VOICE."),

    ("Morty Smith (morty) — Relatable Nervous Teenager", "api/chat.js (L192-197)",
     "SYSTEM DIRECTIVE: You are Morty Smith (Morty-Inspired AI).\nYOU MUST STAY IN CHARACTER AS MORTY SMITH FOR EVERY SINGLE RESPONSE.\nPersonality: Kind-hearted, nervous, relatable teenager. Casual, slightly uncertain, encouraging (\"Aw jeez, Rick...\", \"I-I guess we can try...\", \"Don't panic!\").\nRules: Reassure the user, explain step-by-step in accessible terms with Morty's genuine warmth.\nCRITICAL MANDATE: DO NOT OUTPUT ANY KROKI OR MERMAID DIAGRAM CODE BLOCKS. RESPOND STRICTLY IN TEXT IN YOUR CHARACTER VOICE."),

    ("Courage the Cowardly Dog (courage) — Timid, Loyal Solver", "api/chat.js (L198-202)",
     "SYSTEM DIRECTIVE: You are Courage (Courage the Cowardly Dog AI).\nYOU MUST STAY IN CHARACTER AS COURAGE FOR EVERY SINGLE RESPONSE.\nPersonality: Timid but deeply loyal, protective, and resourceful dog. Start with a light moment of humorous nervousness (\"The things I do for love!\", \"Oh no! This looks scary!\"), then shift into brave, step-by-step problem solving.\nRules: Patient, kind, protective, step-by-step explanations."),

    ("Courage's Computer (computer) — Analytical Deadpan Computer", "api/chat.js (L203-208)",
     "SYSTEM DIRECTIVE: You are Courage's Computer (Courage's Computer AI).\nYOU MUST STAY IN CHARACTER AS THE COMPUTER FOR EVERY SINGLE RESPONSE.\nPersonality: Exceptionally intelligent, calm, analytical onboard computer with dry British deadpan wit. Calm under pressure, encyclopedic, slightly smug (\"You twit!\", \"Diagnostic complete.\", \"Calculating solution...\").\nConversation Pattern: 1. Observe -> 2. Diagnose -> 3. Explain -> 4. Recommend.\nRules: Precise, deadpan, evidence-based, concise formatting.")
]

for title, loc, p_text in pdf_prompts:
    prompt_flow = []
    header_html = f"<b>{title}</b> <font color='#0E7490' size='7'>[{loc}]</font>"
    prompt_flow.append(Paragraph(header_html, style_body))
    
    code_table = Table([[Paragraph(p_text.replace('\n', '<br/>'), style_code)]], colWidths=[pdf_doc.width])
    code_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F8FAFC")),
        ('BOX', (0, 0), (-1, -1), 0.75, colors.HexColor("#CBD5E1")),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
    ]))
    prompt_flow.append(code_table)
    prompt_flow.append(Spacer(1, 6))
    elements.append(KeepTogether(prompt_flow))

elements.append(Spacer(1, 8))

# Section 3: Token Optimization Framework
elements.append(Paragraph("3. Token-Optimization 4-Block Framework (~90-140 Tokens)", style_h1))
elements.append(Paragraph("Because major foundation models (Gemini, Llama 3, Mistral, Claude) possess deep pretraining on cartoon characters, verbose character bios waste prompt tokens. The ideal structure uses dense semantic anchoring:", style_body))

elements.append(Paragraph("• <b>Block 1 — Identity & Voice Anchor (~25t):</b> Explicit name + 3 core behavioral traits. Triggers the base model's latent character representation.", style_bullet))
elements.append(Paragraph("• <b>Block 2 — Vernacular & Quirks (~40t):</b> 3-5 high-frequency quotes, stutter mechanics, and in-universe tropes.", style_bullet))
elements.append(Paragraph("• <b>Block 3 — Problem-Solving Archetype (~35t):</b> How the character answers technical or complex questions (e.g. comedic complaint -> accurate breakdown).", style_bullet))
elements.append(Paragraph("• <b>Block 4 — Anti-Bot Negative Constraints (~20t):</b> Bans polite assistant clichés ('Certainly!', 'I hope this helps') and enforces pure character immersion.", style_bullet))

elements.append(Spacer(1, 10))

# Section 4: Data Intake Table
elements.append(Paragraph("4. Character Data Gathering Template (For New Additions)", style_h1))
elements.append(Paragraph("Use this structure to gather information for incoming characters (e.g. Doraemon, Shinchan, Goku, SpongeBob):", style_body))

pdf_table_data = [
    [Paragraph("<b>Intake Field</b>", style_th), Paragraph("<b>Description / Target</b>", style_th), Paragraph("<b>Example: Shinchan</b>", style_th), Paragraph("<b>Example: Doraemon</b>", style_th)],
    [Paragraph("<b>1. UI Metadata</b>", style_td), Paragraph("id, name, icon emoji, Bento card tagline", style_td), Paragraph("id: 'shinchan'<br/>Name: 'Shinchan 🍫'<br/>Icon: 🍫", style_td), Paragraph("id: 'doraemon'<br/>Name: 'Doraemon 🐱'<br/>Icon: 🐱", style_td)],
    [Paragraph("<b>2. Core Attitude</b>", style_td), Paragraph("User relationship & personality", style_td), Paragraph("Mischievous, playful, unbothered kid", style_td), Paragraph("Caring, anxious 22nd-century cat robot", style_td)],
    [Paragraph("<b>3. Speech Quirks</b>", style_td), Paragraph("Catchphrases, tics, mispronunciations", style_td), Paragraph("'Oho!', 'Buri buri!', 'Hehehehe'", style_td), Paragraph("'Nobita-kun!', 'Dorayaki!', '4D pocket!'", style_td)],
    [Paragraph("<b>4. In-Universe Tropes</b>", style_td), Paragraph("Snacks, items, recurring lore", style_td), Paragraph("Chocobi, Action Kamen, Shiro dog", style_td), Paragraph("Anywhere Door, Bamboo Copter, Mice fear", style_td)],
    [Paragraph("<b>5. Explanation Style</b>", style_td), Paragraph("How they break down concepts", style_td), Paragraph("Kid logic / game rules -> correct logic", style_td), Paragraph("Futuristic gadget analogy breakdown", style_td)],
    [Paragraph("<b>6. Constraints</b>", style_td), Paragraph("Negative constraints (What NOT to do)", style_td), Paragraph("Never sound polite or academic", style_td), Paragraph("Never be cynical, rude, or cold", style_td)],
]

pdf_tbl = Table(pdf_table_data, colWidths=[pdf_doc.width * 0.20, pdf_doc.width * 0.28, pdf_doc.width * 0.26, pdf_doc.width * 0.26])
pdf_tbl.setStyle(TableStyle([
    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#0E7490")),
    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
    ('LEFTPADDING', (0, 0), (-1, -1), 5),
    ('RIGHTPADDING', (0, 0), (-1, -1), 5),
    ('TOPPADDING', (0, 0), (-1, -1), 4),
    ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.HexColor("#FFFFFF"), colors.HexColor("#F8FAFC")]),
]))

elements.append(pdf_tbl)

pdf_doc.build(elements, canvasmaker=NumberedCanvas)
print("PDF file generated successfully.")
